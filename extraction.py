"""
PDF / Word text extraction.

- .docx -> python-docx, paragraph text joined with newlines (tables flattened
  row-by-row, cells joined with " | ").
- .pdf -> pdfplumber per-page text extraction. If a page yields no/near-empty
  text (scanned/image-based page), fall back to OCR for that page only via
  pdf2image -> pytesseract, IF those are installed. If not, that page is
  left as empty text rather than crashing the whole extraction (real
  sample found, tote bag Jul 5 batch: trailing screenshot/confidence-
  summary pages had zero embedded text at all -- unlike every prior sample,
  which always had at least a URL/timestamp footer keeping them above the
  OCR threshold -- and this environment doesn't have pdf2image installed).
  Real pattern content (Materials/Stitch Guide/Pattern Steps/Finishing)
  has always been on text-layer pages in every real sample seen so far;
  OCR is a fallback for genuinely scanned pages, not something that should
  hard-crash QA of an otherwise-fine PDF just because it's unavailable.
"""
import os
import re
import sys

MIN_CHARS_BEFORE_OCR_FALLBACK = 20

# Real OCR character confusion (scarf-mossribbed, Jul 15 batch): the
# lowercase "l" in "Sl st" (slip stitch) gets misread as a capital "I" or
# a pipe "|" -- "SI st", "S| st", "sI st" -- inconsistently across a single
# file, alongside correctly-read "sl st" elsewhere. Narrow substitution:
# only touches "S<confusable> st", not "l"/"I"/"|" generally (which would
# risk corrupting unrelated text).
_RE_OCR_SL_ST_CONFUSION = re.compile(r"\bs[lI|]\s*st\b", re.I)


def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return _extract_docx(path)
    elif ext == ".pdf":
        return _extract_pdf(path)
    else:
        raise ValueError(f"Unsupported file type: {ext} (expected .pdf or .docx)")


def _extract_docx(path: str) -> str:
    import docx

    d = docx.Document(path)
    lines = []
    for para in d.paragraphs:
        lines.append(para.text)
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_pdf(path: str) -> str:
    import pdfplumber

    pages_text = []
    ocr_pages = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if len(text.strip()) < MIN_CHARS_BEFORE_OCR_FALLBACK:
                ocr_pages.append(i)
                pages_text.append(None)  # placeholder, filled in below
            else:
                pages_text.append(text)

    if ocr_pages:
        # Real sample (sweater, Jul 12 batch; scarf, Jul 15 batch): a new
        # LoopDreams export template renders every page as vector graphics
        # with no embedded text at all, so EVERY page needs OCR -- at
        # ~15-20s/page, a batch of these can run for many minutes with zero
        # visible output otherwise, easily looking hung. This note (and the
        # per-page one in _ocr_pages below) give visible progress instead.
        print(
            f"Note: {os.path.basename(path)} has {len(ocr_pages)} page(s) with no text layer -- "
            f"OCR'ing them now (this can take a while for large files).",
            file=sys.stderr,
        )
        ocr_results = _ocr_pages(path, ocr_pages)
        for idx, text in zip(ocr_pages, ocr_results):
            pages_text[idx] = text

    return "\n".join(t or "" for t in pages_text)


def _ocr_pages(path: str, page_indices) -> list:
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError as e:
        print(
            f"Warning: OCR fallback unavailable ({e}) -- {len(page_indices)} low-text page(s) "
            f"will be treated as empty rather than crashing extraction. Install pdf2image + "
            f"pytesseract (+ system tesseract) to OCR scanned/image-only pages.",
            file=sys.stderr,
        )
        return ["" for _ in page_indices]

    results = []
    total = len(page_indices)
    for n, idx in enumerate(page_indices, start=1):
        print(f"  OCR'ing page {idx + 1} ({n}/{total})...", file=sys.stderr, flush=True)
        images = convert_from_path(path, first_page=idx + 1, last_page=idx + 1, dpi=300)
        if images:
            # Real bug (scarf-mossribbed, Jul 15 batch): Tesseract's default
            # page segmentation (PSM 3, "fully automatic") badly scrambled
            # the reading order on a page with many short, densely-stacked
            # "Row N" badge+instruction pairs (a ribbing section) --
            # producing all the row-number badges first, then all the
            # instruction text second, completely decoupled from their real
            # row numbers. PSM 6 ("assume a single uniform block of text")
            # reads this correctly, and was verified to still correctly
            # handle a genuine multi-column layout (the Pattern Overview's
            # Finished Size/Yarn/Hook/Yardage table) -- not just narrowly
            # fixing this one page at the cost of breaking others.
            raw = pytesseract.image_to_string(images[0], config="--psm 6")
            results.append(_RE_OCR_SL_ST_CONFUSION.sub("sl st", raw))
        else:
            results.append("")
    return results
